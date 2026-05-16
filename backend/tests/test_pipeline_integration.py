import tempfile
import os
import sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))


def test_word_pipeline_end_to_end():
    try:
        from docx import Document
    except ImportError:
        return

    with tempfile.TemporaryDirectory() as tmpdir:
        os.environ['KB_DATA_DIR'] = tmpdir
        from db.models import init_db
        from pipeline.word import WordPipeline
        from card.store import CardStore
        import sqlite3

        conn = sqlite3.connect(os.path.join(tmpdir, "test.db"))
        conn.row_factory = sqlite3.Row
        init_db(conn)
        conn.close()

        doc_path = os.path.join(tmpdir, "test.docx")
        doc = Document()
        doc.add_heading("Overview", level=1)
        doc.add_paragraph("AE800 is a 4K video conference system with advanced features.")
        doc.add_heading("Pricing", level=2)
        doc.add_paragraph("AE800 price: 138000 RMB per unit.")
        doc.save(doc_path)

        pipeline = WordPipeline(data_dir=tmpdir)
        cards = pipeline.parse(doc_path)
        assert len(cards) >= 2

        store = CardStore(data_dir=tmpdir)
        store.save_batch(cards)

        loaded = store.get(cards[0]["id"])
        assert loaded is not None
        assert loaded["source_type"] == "word"

        result = store.list_cards(source_type="word")
        assert result["total"] >= 2


def test_markdown_pipeline_end_to_end():
    with tempfile.TemporaryDirectory() as tmpdir:
        from pipeline.markdown import MarkdownPipeline
        from card.store import CardStore

        md_path = os.path.join(tmpdir, "solution.md")
        with open(md_path, "w") as f:
            f.write("# 公安方案\n\n巡查督导方案内容。\n\n## 教育方案\n\n远程课堂内容。")

        pipeline = MarkdownPipeline(data_dir=tmpdir)
        cards = pipeline.parse(md_path)
        assert len(cards) >= 2

        store = CardStore(data_dir=tmpdir)
        store.save_batch(cards)
        assert store.list_cards()["total"] >= 2


def test_txt_pipeline_end_to_end():
    with tempfile.TemporaryDirectory() as tmpdir:
        from pipeline.txt import TxtPipeline
        from card.store import CardStore

        txt_path = os.path.join(tmpdir, "cases.txt")
        with open(txt_path, "w") as f:
            f.write("案例一\n\n案例内容。\n\n案例二\n\n更多内容。")

        pipeline = TxtPipeline(data_dir=tmpdir)
        cards = pipeline.parse(txt_path)
        assert len(cards) >= 2

        store = CardStore(data_dir=tmpdir)
        store.save_batch(cards)
        assert store.list_cards()["total"] >= 2
