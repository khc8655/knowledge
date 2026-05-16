import tempfile
import os
import sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from pipeline.word import WordPipeline


def test_parse_simple_docx():
    try:
        from docx import Document
    except ImportError:
        return

    with tempfile.TemporaryDirectory() as tmpdir:
        doc_path = os.path.join(tmpdir, "test.docx")
        doc = Document()
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("This is the introduction paragraph with enough content to test.")
        doc.add_heading("Features", level=2)
        doc.add_paragraph("Feature A description here.")
        doc.add_paragraph("Feature B description here.")
        doc.save(doc_path)

        pipeline = WordPipeline(data_dir=tmpdir)
        cards = pipeline.parse(doc_path)

        assert len(cards) >= 2
        assert cards[0]["source_type"] == "word"
        assert cards[0]["level"] >= 1
